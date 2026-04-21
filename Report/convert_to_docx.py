#!/usr/bin/env python3
"""
convert_to_docx.py — Convert Vista_Report.md → Vista_Report.docx

Fixes applied
-------------
* Title page content no longer duplicated (div section stripped before parse)
* Figure captions detected for both *Fig. and **Fig. formats
* Figure captions rendered BELOW diagrams (academic convention)
* Table captions rendered ABOVE tables (academic convention)
* Chapter labels (# CHAPTER I) rendered as centred label + page break
* Chapter titles (## Chapter N:) rendered with prominence equal to H1
* Page numbers added to footer
* Inline math $...$ italicised
* All 23 Mermaid diagrams rendered via mmdc as high-res PNG
* Block math $$...$$ rendered via matplotlib as image

Usage (from the Report/ folder with venv active)
-------------------------------------------------
    source ../vista_env/bin/activate
    python3 convert_to_docx.py
"""
from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Paths ─────────────────────────────────────────────────────────────────────
HERE     = Path(__file__).parent
MD_PATH  = HERE / "Vista_Report.md"
OUT_PATH = HERE / "Vista_Report.docx"
MMDC_BIN = HERE / "node_modules" / ".bin" / "mmdc"
PUPPET   = HERE / "puppeteer-config.json"

# ── Brand colours ─────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1E, 0x3A, 0x5F)
TEAL  = RGBColor(0x0F, 0x76, 0x6E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
RED   = RGBColor(0xC7, 0x25, 0x4E)

# ── XML / style helpers ────────────────────────────────────────────────────────
def _shd(elem, fill: str) -> None:
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),   "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"),  fill)
    elem.append(s)


def _shade_cell(cell, fill: str) -> None:
    _shd(cell._tc.get_or_add_tcPr(), fill)


def _cell_border(cell, color: str = "B0BEC5") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    bdr   = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        bdr.append(el)
    tc_pr.append(bdr)


def _hr(doc: Document) -> None:
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1E3A5F")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ── Document page setup ────────────────────────────────────────────────────────
def _setup(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_height   = Inches(11)
    sec.page_width    = Inches(8.5)
    sec.top_margin    = Inches(1.5)
    sec.bottom_margin = Inches(1.5)
    sec.left_margin   = Inches(2.0)
    sec.right_margin  = Inches(1.33)

    normal = doc.styles["Normal"]
    normal.font.name  = "Times New Roman"
    normal.font.size  = Pt(12)
    normal.font.color.rgb = DARK
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = Pt(18)
    pf.space_after       = Pt(6)
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_page_numbers(doc: Document) -> None:
    """Add centred page numbers to the footer of the first section."""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = DARK

    for fld_type, text in [("begin", ""), ("separate", ""), ("end", "")]:
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), fld_type)
        run._r.append(fld)

    instr_run = para.add_run()
    instr_run.font.name = "Times New Roman"
    instr_run.font.size = Pt(11)
    i_el = OxmlElement("w:instrText")
    i_el.set(qn("xml:space"), "preserve")
    i_el.text = " PAGE "
    instr_run._r.append(i_el)

    end_run = para.add_run()
    end_run.font.name = "Times New Roman"
    end_run.font.size = Pt(11)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


# ── Mermaid rendering ─────────────────────────────────────────────────────────
def _render_mermaid(mermaid_src: str, out_dir: Path, idx: int) -> Path | None:
    """Render a Mermaid diagram to PNG; returns PNG path or None on error."""
    if not MMDC_BIN.exists():
        return None
    src_file = out_dir / f"diagram_{idx:03d}.mmd"
    png_file = out_dir / f"diagram_{idx:03d}.png"
    src_file.write_text(mermaid_src, encoding="utf-8")
    cmd = [
        str(MMDC_BIN),
        "-i", str(src_file),
        "-o", str(png_file),
        "-w", "2000",
        "-H", "1400",
        "-s", "2",
        "--backgroundColor", "white",
    ]
    if PUPPET.exists():
        cmd += ["-p", str(PUPPET)]
    try:
        result = subprocess.run(cmd, capture_output=True, timeout=90, check=True)
        if png_file.exists() and png_file.stat().st_size > 100:
            print(f"  [mmdc] diagram_{idx:03d} OK ({png_file.stat().st_size // 1024} KB)")
            return png_file
        print(f"  [mmdc] diagram_{idx:03d} produced empty/missing PNG")
    except subprocess.CalledProcessError as e:
        stderr = e.stderr.decode(errors="replace")[:400] if e.stderr else ""
        print(f"  [mmdc] diagram_{idx:03d} failed (exit {e.returncode}): {stderr}")
    except Exception as e:
        print(f"  [mmdc] diagram_{idx:03d} error: {e}")
    return None


# ── Math formula image rendering ───────────────────────────────────────────────
_math_img_counter = 0


def _render_math_img(latex_src: str, out_dir: Path, idx: int) -> Path | None:
    """Render a LaTeX block expression to PNG via matplotlib."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        png_file = out_dir / f"math_{idx:03d}.png"
        fig = plt.figure(figsize=(7, 1.1))
        fig.patch.set_facecolor("white")
        fig.text(0.5, 0.5, f"${latex_src}$",
                 ha="center", va="center",
                 fontsize=15, color="#1e293b")
        plt.savefig(str(png_file), dpi=220, bbox_inches="tight",
                    facecolor="white", edgecolor="none", pad_inches=0.15)
        plt.close(fig)
        if png_file.exists() and png_file.stat().st_size > 100:
            return png_file
    except Exception as e:
        print(f"  [math] math_{idx:03d} failed: {e}")
    return None


def _math_block(doc: Document, text: str, tmp: Path | None = None) -> None:
    """Render a block math expression as an image (preferred) or plain text."""
    global _math_img_counter
    latex_src = re.sub(r"\$\$", "", text).strip()
    rendered = False
    if tmp is not None:
        _math_img_counter += 1
        png = _render_math_img(latex_src, tmp, _math_img_counter)
        if png:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(10)
            p.add_run().add_picture(str(png), width=Inches(4.5))
            rendered = True
    if not rendered:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(latex_src)
        r.font.name = "Cambria Math"
        r.font.size = Pt(12)
        r.italic    = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)


# ── Image embedding ────────────────────────────────────────────────────────────
def _embed_image(doc: Document, png: Path, caption: str = "") -> None:
    """Insert a PNG as a centred figure with caption BELOW (academic convention)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    p.add_run().add_picture(str(png), width=Inches(5.0))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(12)
        cr = cp.add_run(caption)
        cr.italic          = True
        cr.font.size       = Pt(10)
        cr.font.name       = "Times New Roman"
        cr.font.color.rgb  = TEAL


# ── Title page ─────────────────────────────────────────────────────────────────
def _title_page(doc: Document) -> None:
    def cline(text: str, bold=False, size=12, color=NAVY,
              sb=0, sa=8, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        r = p.add_run(text)
        r.bold = bold; r.font.name = "Times New Roman"
        r.font.size = Pt(size); r.font.color.rgb = color

    cline("GHULAM ISHAQ KHAN INSTITUTE OF", bold=True, size=14, sb=0, sa=2)
    cline("ENGINEERING SCIENCES AND TECHNOLOGY", bold=True, size=14, sb=0, sa=4)
    cline("Faculty of Computer Engineering", size=12, sb=0, sa=30)

    cline("VISTA", bold=True, size=24, sb=20, sa=4)
    cline("Verification using Intelligent Systems for Test Automation",
          bold=True, size=14, sb=0, sa=4)
    cline("Senior Design Project (SDP) — Final Year Project Report",
          size=12, sb=0, sa=30)

    cline("Submitted by:", bold=True, size=12, sb=0, sa=8)

    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Name", "Registration Number"]):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(11)
        r.font.color.rgb = WHITE
        _shade_cell(cell, "1E3A5F")
        _cell_border(cell)
    for ri, (nm, reg) in enumerate([
            ("Asad Ali",     "2021-CE-XXX"),
            ("Sheikh Abbas", "2021-CE-XXX"),
            ("Mian Arqam",   "2021-CE-XXX")], start=1):
        for ci, val in enumerate([nm, reg]):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(val)
            r.font.name = "Times New Roman"; r.font.size = Pt(11)
            _cell_border(cell)
            if ri % 2 == 0:
                _shade_cell(cell, "EEF3FB")
    doc.add_paragraph()

    cline("Advisor: Dr. Fahad Bin Muslim",              size=12, sb=12, sa=2)
    cline("Co-Advisor: Dr. Waqar",                      size=12, sb=0,  sa=2)
    cline("Industrial Partner: NECOP",                   size=12, sb=0,  sa=2)
    cline("Industrial Mentors: Dr. Haroon  |  Dr. Majeed", size=12, sb=0, sa=20)
    cline("Degree: BS Computer Engineering",             size=11, sb=0,  sa=2)
    cline("Session: 2021 – 2025",                        size=11, sb=0,  sa=2)
    cline("April 2026",                                  size=11, sb=0,  sa=0)
    doc.add_page_break()


# ── Inline formatting ──────────────────────────────────────────────────────────
_INLINE = re.compile(
    r"\*\*\*(?P<boit>[^*]+?)\*\*\*"
    r"|\*\*(?P<bo>[^*]+?)\*\*"
    r"|\*(?P<it>[^*]+?)\*"
    r"|`(?P<code>[^`]+?)`"
)


def _clean_text(text: str) -> str:
    """Strip HTML entities and tags, convert inline math to italic markers."""
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\\([|`*_])", r"\1", text)
    # Convert inline $math$ to *math* (italic) stripping LaTeX commands
    def _math_to_italic(m: re.Match) -> str:
        inner = m.group(1)
        inner = re.sub(r"\\text\{([^}]+)\}", r"\1", inner)
        inner = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"(\1)/(\2)", inner)
        inner = re.sub(r"\\left|\\right|\\!", "", inner)
        inner = re.sub(r"\\[a-zA-Z]+", "", inner)
        inner = re.sub(r"[{}]", "", inner)
        inner = inner.strip()
        return f"*{inner}*" if inner else ""
    text = re.sub(r"\$([^$\n]+)\$", _math_to_italic, text)
    return text


def _add_inline(para, text: str, size=12, bold=False,
                italic=False, color=DARK) -> None:
    text = _clean_text(text)
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            r.font.name = "Times New Roman"; r.font.size = Pt(size)
            r.bold = bold; r.italic = italic; r.font.color.rgb = color
        pos = m.end()
        boit = m.group("boit"); bo = m.group("bo")
        it   = m.group("it");   code = m.group("code")
        if boit:
            r = para.add_run(boit)
            r.bold = True; r.italic = True
            r.font.name = "Times New Roman"; r.font.size = Pt(size)
            r.font.color.rgb = color
        elif bo:
            r = para.add_run(bo)
            r.bold = True; r.italic = italic
            r.font.name = "Times New Roman"; r.font.size = Pt(size)
            r.font.color.rgb = color
        elif it:
            r = para.add_run(it)
            r.bold = bold; r.italic = True
            r.font.name = "Times New Roman"; r.font.size = Pt(size)
            r.font.color.rgb = color
        elif code:
            r = para.add_run(code)
            r.font.name = "Courier New"; r.font.size = Pt(10)
            r.font.color.rgb = RED
    if pos < len(text):
        r = para.add_run(text[pos:])
        r.font.name = "Times New Roman"; r.font.size = Pt(size)
        r.bold = bold; r.italic = italic; r.font.color.rgb = color


# ── Headings ──────────────────────────────────────────────────────────────────
def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    cfg = {
        1: (16, NAVY, True,  False, Pt(16), Pt(10)),
        2: (14, NAVY, True,  False, Pt(12), Pt(8)),
        3: (12, TEAL, True,  False, Pt(8),  Pt(6)),
        4: (11, DARK, True,  True,  Pt(6),  Pt(4)),
    }
    sz, col, bd, it, sb, sa = cfg.get(level, cfg[4])
    p.paragraph_format.space_before = sb
    p.paragraph_format.space_after  = sa
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.upper() if level == 1 else text)
    r.bold = bd; r.italic = it
    r.font.name = "Times New Roman"; r.font.size = Pt(sz)
    r.font.color.rgb = col


def _chapter_label(doc: Document, label: str) -> None:
    """Render 'CHAPTER I' as a centred chapter label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(label)
    r.bold = True
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(14)
    r.font.color.rgb = NAVY


# ── Table ──────────────────────────────────────────────────────────────────────
def _render_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri]
        for ci in range(ncols):
            cell = row.cells[ci]
            val  = row_data[ci].strip() if ci < len(row_data) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline(para, val, size=10,
                        bold=(ri == 0), color=WHITE if ri == 0 else DARK)
            _cell_border(cell, "AAAAAA")
            if ri == 0:
                _shade_cell(cell, "1E3A5F")
            elif ri % 2 == 0:
                _shade_cell(cell, "EEF3FB")
    doc.add_paragraph()


# ── Code block ─────────────────────────────────────────────────────────────────
def _render_code(doc: Document, lines: list[str]) -> None:
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    _shade_cell(cell, "1E293B")
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(line)
        r.font.name = "Courier New"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()


# ── List item ──────────────────────────────────────────────────────────────────
def _list_item(doc: Document, text: str, level: int = 0,
               numbered: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent       = Cm(0.6 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after       = Pt(3)
    bullet = ("•" if level == 0 else "–") if not numbered else f"{level + 1}."
    r = p.add_run(f"{bullet}  ")
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    r.bold = True; r.font.color.rgb = NAVY
    _add_inline(p, text)


# ── Caption (above tables / below figures via _embed_image) ───────────────────
def _caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.strip("*_ "))
    r.italic = True; r.font.name = "Times New Roman"
    r.font.size = Pt(10); r.font.color.rgb = TEAL


# ── Body paragraph ─────────────────────────────────────────────────────────────
def _body(doc: Document, text: str,
          align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    _add_inline(p, text)


# ── Main parser ────────────────────────────────────────────────────────────────
def _parse(doc: Document, md: str, tmp: Path) -> None:  # noqa: C901
    lines    = md.splitlines()
    n        = len(lines)
    i        = 0
    diag_n   = 0

    code_buf:  list[str] = []
    code_lang  = ""
    in_code    = False

    table_buf: list[list[str]] = []
    in_table   = False

    in_math   = False
    math_buf: list[str] = []

    pending_fig_caption = ""    # stored and rendered BELOW the next figure

    def flush_table() -> None:
        nonlocal in_table, table_buf
        if table_buf:
            _render_table(doc, table_buf)
        in_table = False
        table_buf = []

    # ── Caption detection (matches *Fig.**, **Fig.**, *Table X*, **Table X**) ──
    _CAP_RE = re.compile(
        r"^\*+\s*"
        r"((?:Fig(?:ure)?\.?|Table)\s[\d.]+\s*[—–\-]?\s*.*?)"
        r"\*+\s*$",
        re.IGNORECASE,
    )

    while i < n:
        raw      = lines[i]
        line     = raw.rstrip()
        stripped = line.strip()

        # ── HTML skip ─────────────────────────────────────────────────────────
        if re.match(r"^<(div|/div|br)[\s/>]?", stripped, re.I) or stripped in ("<div>", "</div>", "<br>", "<br/>"):
            i += 1; continue

        # ── Fenced code / mermaid ─────────────────────────────────────────────
        fence = re.match(r"^```(\w*)", stripped)
        if fence and not in_code:
            if in_table:
                flush_table()
            in_code   = True
            code_lang = fence.group(1).lower()
            code_buf  = []
            i += 1; continue

        if stripped == "```" and in_code:
            if code_lang == "mermaid":
                diag_n += 1
                src = "\n".join(code_buf)
                png = _render_mermaid(src, tmp, diag_n)
                if png:
                    _embed_image(doc, png, pending_fig_caption)
                else:
                    _render_code(doc, code_buf)
                    if pending_fig_caption:
                        _caption(doc, pending_fig_caption)
                pending_fig_caption = ""
            else:
                _render_code(doc, code_buf)
            in_code = False; code_buf = []; code_lang = ""
            i += 1; continue

        if in_code:
            code_buf.append(raw.rstrip())
            i += 1; continue

        # ── Block math ($$...$$) ──────────────────────────────────────────────
        if stripped.startswith("$$"):
            # Single-line: $$...$$ all on one line
            if stripped.endswith("$$") and stripped != "$$":
                _math_block(doc, stripped, tmp)
            # Multi-line closing $$
            elif in_math:
                math_buf.append(stripped)
                _math_block(doc, "\n".join(math_buf), tmp)
                in_math = False; math_buf = []
            # Multi-line opening $$
            else:
                in_math = True; math_buf = [stripped]
            i += 1; continue
        if in_math:
            math_buf.append(stripped); i += 1; continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^---+\s*$", stripped):
            if in_table:
                flush_table()
            i += 1; continue   # front-matter separators — skip silently

        # ── Chapter page break markers (# CHAPTER BREAK inserted by main) ─────
        if stripped == "# CHAPTER BREAK":
            if in_table:
                flush_table()
            doc.add_page_break()
            i += 1; continue

        # ── Chapter labels (# CHAPTER I / # CHAPTER II etc.) ─────────────────
        chap_lbl = re.match(r"^#\s+(CHAPTER\s+\S+.*)", stripped)
        if chap_lbl:
            if in_table:
                flush_table()
            doc.add_page_break()
            _chapter_label(doc, chap_lbl.group(1).strip())
            i += 1; continue

        # ── Headings ──────────────────────────────────────────────────────────
        h = re.match(r"^(#{1,4})\s+(.+)", stripped)
        if h:
            if in_table:
                flush_table()
            level = len(h.group(1))
            text  = h.group(2).strip()

            # Chapter titles (## Chapter N: ...) — strip prefix, render at level-1
            if level == 2 and re.match(r"Chapter\s+\d+", text, re.I):
                title_only = re.sub(r"^Chapter\s+\d+[:.]\s*", "", text, flags=re.I).strip()
                _heading(doc, title_only or text, level=1)
            else:
                _heading(doc, text, level)
            i += 1; continue

        # ── Pipe tables ───────────────────────────────────────────────────────
        if stripped.startswith("|"):
            # Separator rows: skip
            clean = stripped.replace(" ", "")
            if re.match(r"^\|[\-:|]+(\|[\-:|]*)*\|?$", clean):
                i += 1; continue
            cols = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True; table_buf = []
            table_buf.append(cols)
            i += 1; continue
        else:
            if in_table:
                flush_table()

        # ── Blank lines ───────────────────────────────────────────────────────
        if not stripped:
            i += 1; continue

        # ── Caption detection (*Fig. / **Fig. / *Table X / **Table X) ─────────
        cap_m = _CAP_RE.match(stripped)
        if cap_m:
            cap_text = cap_m.group(1).strip().rstrip("*").strip()
            is_table_cap = bool(re.match(r"Table\s+[\d.]", cap_text, re.I))
            if is_table_cap:
                # Table captions go ABOVE the table — render immediately
                _caption(doc, cap_text)
                pending_fig_caption = ""
            else:
                # Figure captions go BELOW the figure — defer to _embed_image
                pending_fig_caption = cap_text
            i += 1; continue

        # ── Bullet lists ──────────────────────────────────────────────────────
        bul = re.match(r"^(\s*)[*\-]\s+(.+)", line)
        if bul:
            _list_item(doc, bul.group(2), level=len(bul.group(1)) // 2)
            i += 1; continue

        # ── Numbered lists ────────────────────────────────────────────────────
        num = re.match(r"^(\s*)\d+\.\s+(.+)", line)
        if num:
            _list_item(doc, num.group(2), level=len(num.group(1)) // 2,
                       numbered=True)
            i += 1; continue

        # ── Reference list [N] entries ────────────────────────────────────────
        ref = re.match(r"^\[(\d+)\]\s+(.+)", stripped)
        if ref:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent       = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            p.paragraph_format.space_after       = Pt(4)
            r_num = p.add_run(f"[{ref.group(1)}]  ")
            r_num.bold = True; r_num.font.name = "Times New Roman"
            r_num.font.size = Pt(11); r_num.font.color.rgb = NAVY
            _add_inline(p, ref.group(2), size=11)
            i += 1; continue

        # ── Plain paragraph ───────────────────────────────────────────────────
        _body(doc, stripped)
        i += 1

    if in_table:
        flush_table()


# ── Entry point ────────────────────────────────────────────────────────────────
def main() -> None:
    if not MMDC_BIN.exists():
        print(f"WARNING: mmdc not found at {MMDC_BIN}")
        print("         Run: npm install   in the Report/ folder first.")
        print("         Diagrams will fall back to text blocks.\n")
    else:
        print(f"mmdc     : {MMDC_BIN}")

    print(f"Reading  : {MD_PATH}")
    md = MD_PATH.read_text(encoding="utf-8")

    # ── Strip the <div>...</div> title page block (handled by _title_page()) ──
    div_end = md.find("</div>")
    if div_end >= 0:
        md = md[div_end + 6:].lstrip()
        # Also drop the first bare --- separator that follows
        if md.startswith("---"):
            md = md[3:].lstrip()

    # ── Inject page breaks before front-matter sections ───────────────────────
    _PAGE_BREAKS = [
        "Certificate of Approval", "SDGs Inclusion", "Abstract",
        "Acknowledgements", "Table of Contents", "List of Figures",
        "List of Tables", "Glossary", "References",
        "Appendix A", "Appendix B", "Appendix C",
    ]
    for sec in _PAGE_BREAKS:
        md = re.sub(
            rf"(^## {re.escape(sec)}\b)",
            r"\n# CHAPTER BREAK\n\g<1>",
            md,
            flags=re.MULTILINE,
        )

    doc = Document()
    _setup(doc)
    _title_page(doc)
    _add_page_numbers(doc)

    with tempfile.TemporaryDirectory(prefix="vista_diagrams_") as tmp_s:
        tmp = Path(tmp_s)
        print(f"Temp dir : {tmp}")
        _parse(doc, md, tmp)

    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size // 1024
    print(f"\nSaved    : {OUT_PATH}  ({size_kb} KB)")
    print("Done.")


if __name__ == "__main__":
    main()
